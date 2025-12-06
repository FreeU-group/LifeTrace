"""工作区文件管理相关路由"""

import os
import shutil
from datetime import datetime
from pathlib import Path

from fastapi import APIRouter, File, HTTPException, Query, UploadFile
from fastapi.responses import StreamingResponse

from lifetrace.routers import dependencies as deps
from lifetrace.schemas.workspace import (
    CreateFileRequest,
    CreateFileResponse,
    CreateFolderRequest,
    CreateFolderResponse,
    CreateWorkspaceProjectRequest,
    CreateWorkspaceProjectResponse,
    DeleteFileRequest,
    DeleteFileResponse,
    DeleteWorkspaceProjectRequest,
    DeleteWorkspaceProjectResponse,
    DocumentAction,
    DocumentAIRequest,
    DocumentAIResponse,
    FileContentResponse,
    FileNode,
    GenerateSlidesRequest,
    GenerateSlidesResponse,
    GeneratedSlideInfo,
    PreviewSlidesPromptsResponse,
    ProjectType,
    RenameFileRequest,
    RenameFileResponse,
    RenameWorkspaceProjectRequest,
    RenameWorkspaceProjectResponse,
    SaveFileRequest,
    SaveFileResponse,
    SlideImageInfo,
    SlidesImagesResponse,
    UploadFileResponse,
    UploadImageResponse,
    WorkspaceFilesResponse,
    WorkspaceProject,
    WorkspaceProjectsResponse,
)
from lifetrace.util.logging_config import get_logger
from lifetrace.util.prompt_loader import get_prompt
from lifetrace.util.token_usage_logger import log_token_usage

logger = get_logger()

router = APIRouter(prefix="/api/workspace", tags=["workspace"])


# ==================== 项目管理 API ====================


def get_project_info(project_path: Path) -> dict:
    """获取项目信息（文件数、最后修改时间等）"""
    file_count = 0
    last_modified = None

    for item in project_path.rglob("*"):
        if item.is_file() and not item.name.startswith("."):
            file_count += 1
            mtime = datetime.fromtimestamp(item.stat().st_mtime)
            if last_modified is None or mtime > last_modified:
                last_modified = mtime

    # 如果没有文件，使用文件夹的修改时间
    if last_modified is None:
        last_modified = datetime.fromtimestamp(project_path.stat().st_mtime)

    return {
        "file_count": file_count,
        "last_modified": last_modified.isoformat() if last_modified else None,
    }


@router.get("/projects", response_model=WorkspaceProjectsResponse)
async def get_workspace_projects():
    """获取工作区项目列表（一级文件夹）"""
    try:
        workspace_dir = deps.config.workspace_dir

        # 确保目录存在
        if not os.path.exists(workspace_dir):
            os.makedirs(workspace_dir, exist_ok=True)
            logger.info(f"创建工作区目录: {workspace_dir}")
            return WorkspaceProjectsResponse(projects=[], total=0)

        projects = []
        root = Path(workspace_dir)

        # 获取所有一级文件夹
        for item in sorted(root.iterdir(), key=lambda x: x.name.lower()):
            if item.is_dir() and not item.name.startswith("."):
                info = get_project_info(item)
                project = WorkspaceProject(
                    id=item.name,
                    name=item.name,
                    file_count=info["file_count"],
                    last_modified=info["last_modified"],
                    created_at=datetime.fromtimestamp(item.stat().st_ctime).isoformat(),
                )
                projects.append(project)

        logger.debug(f"获取工作区项目列表成功，共 {len(projects)} 个项目")
        return WorkspaceProjectsResponse(projects=projects, total=len(projects))

    except Exception as e:
        logger.error(f"获取工作区项目列表失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


def _get_outline_template(project_type: ProjectType) -> str:
    """根据项目类型获取大纲模板"""
    type_to_prompt_key = {
        ProjectType.LIBERAL_ARTS: "liberal_arts",
        ProjectType.SCIENCE: "science",
        ProjectType.ENGINEERING: "engineering",
        ProjectType.OTHER: "other",
    }
    prompt_key = type_to_prompt_key.get(project_type, "other")
    return get_prompt("project_outline", prompt_key)


def _get_project_type_name(project_type: ProjectType) -> str:
    """获取项目类型的中文名称"""
    type_names = {
        ProjectType.LIBERAL_ARTS: "文科",
        ProjectType.SCIENCE: "理科",
        ProjectType.ENGINEERING: "工科",
        ProjectType.OTHER: "其他",
    }
    return type_names.get(project_type, "其他")


def _generate_outline_template(project_name: str, project_type: ProjectType) -> str:
    """生成项目大纲模板（仅替换项目名称，不调用 LLM）"""
    template = _get_outline_template(project_type)
    return template.replace("{project_name}", project_name)


def _create_outline_stream_generator(
    project_name: str, project_type: ProjectType, outline_path: Path
):
    """创建流式大纲生成器"""

    def generator():
        template = _get_outline_template(project_type)
        type_name = _get_project_type_name(project_type)

        # 检查 LLM 是否可用
        if not deps.rag_service.llm_client.is_available():
            yield "[错误] LLM 服务不可用"
            return

        try:
            # 获取智能填充的 prompt
            system_prompt = get_prompt("project_outline", "smart_fill_system")
            user_template = get_prompt("project_outline", "smart_fill_user")

            user_message = (
                user_template.replace("{project_name}", project_name)
                .replace("{project_type}", type_name)
                .replace("{template}", template)
            )

            # 流式调用 LLM
            response = deps.rag_service.llm_client.client.chat.completions.create(
                model=deps.rag_service.llm_client.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message},
                ],
                temperature=0.7,
                stream=True,
                stream_options={"include_usage": True},
            )

            full_content = ""
            usage_info = None

            for chunk in response:
                if hasattr(chunk, "usage") and chunk.usage:
                    usage_info = chunk.usage

                if chunk.choices and len(chunk.choices) > 0 and chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    full_content += content
                    yield content

            # 保存完整内容到文件
            outline_path.write_text(full_content, encoding="utf-8")
            logger.info(f"流式生成大纲完成: {project_name}")

            # 记录 token 使用量
            if usage_info:
                log_token_usage(
                    model=deps.rag_service.llm_client.model,
                    input_tokens=usage_info.prompt_tokens,
                    output_tokens=usage_info.completion_tokens,
                    endpoint="workspace_generate_outline_stream",
                    user_query=project_name,
                    response_type="stream",
                    feature_type="workspace_assistant",
                    additional_info={
                        "project_name": project_name,
                        "project_type": project_type.value,
                    },
                )

        except Exception as e:
            logger.error(f"流式生成大纲失败: {e}")
            yield f"\n[错误] 生成失败: {e}"

    return generator


@router.post("/projects/{project_id}/outline/generate")
async def generate_project_outline_stream(
    project_id: str,
    project_type: str = Query("other", description="项目类型"),
):
    """流式生成项目大纲内容

    Args:
        project_id: 项目 ID
        project_type: 项目类型 (liberal_arts, science, engineering, other)
    """
    try:
        workspace_dir = deps.config.workspace_dir
        project_path = Path(workspace_dir) / project_id
        outline_path = project_path / "outline.md"

        # 检查项目是否存在
        if not project_path.exists():

            async def error_gen():
                yield f"[错误] 项目不存在: {project_id}"

            return StreamingResponse(error_gen(), media_type="text/plain; charset=utf-8")

        # 转换项目类型
        try:
            p_type = ProjectType(project_type)
        except ValueError:
            p_type = ProjectType.OTHER

        # 创建流式生成器
        generator = _create_outline_stream_generator(project_id, p_type, outline_path)
        headers = {"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}

        return StreamingResponse(
            generator(), media_type="text/plain; charset=utf-8", headers=headers
        )

    except Exception as e:
        logger.error(f"流式生成大纲失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


def _parse_outline_chapters(outline_content: str) -> list[dict]:
    """解析大纲内容，提取章节信息

    Args:
        outline_content: 大纲内容

    Returns:
        章节列表，每个章节包含 title, points, content
        - title: 章节标题
        - points: 要点列表（- 开头的行）
        - content: 章节的完整内容（标题之间的所有文本）
    """
    chapters = []
    lines = outline_content.strip().split("\n")
    current_chapter = None
    current_points = []
    current_content_lines = []  # 存储当前章节的所有内容行

    for line in lines:
        stripped = line.strip()
        
        # 检测二级标题 (## xxx)
        if stripped.startswith("## "):
            # 保存之前的章节
            if current_chapter:
                # 合并内容行，去除首尾空行
                content_text = "\n".join(current_content_lines).strip()
                chapters.append(
                    {
                        "title": current_chapter,
                        "points": current_points,
                        "content": content_text,
                    }
                )
            # 开始新章节
            current_chapter = stripped[3:].strip()
            current_points = []
            current_content_lines = []
        elif current_chapter:
            # 当前章节下的所有内容
            if stripped:
                # 检测要点 (- xxx)
                if stripped.startswith("- "):
                    point = stripped[2:].strip()
                    if point:
                        current_points.append(point)
                # 所有非空行都添加到内容中
                current_content_lines.append(line)
            else:
                # 保留空行以保持格式
                current_content_lines.append("")

    # 保存最后一个章节
    if current_chapter:
        content_text = "\n".join(current_content_lines).strip()
        chapters.append(
            {
                "title": current_chapter,
                "points": current_points,
                "content": content_text,
            }
        )

    # 过滤掉参考文献章节
    chapters = [
        ch
        for ch in chapters
        if "参考文献" not in ch["title"] and "references" not in ch["title"].lower()
    ]

    return chapters


def _sanitize_filename(title: str) -> str:
    """将章节标题转换为合法的文件名"""
    # 移除中文数字前缀（如 "一、"、"二、" 等）
    import re

    title = re.sub(r"^[一二三四五六七八九十]+、\s*", "", title)
    # 移除数字前缀（如 "1."、"2." 等）
    title = re.sub(r"^\d+\.\s*", "", title)
    # 替换非法字符
    invalid_chars = ["/", "\\", ":", "*", "?", '"', "<", ">", "|", " "]
    for char in invalid_chars:
        title = title.replace(char, "_")
    return title[:50]  # 限制长度


def _create_chapters_stream_generator(project_id: str, outline_content: str, project_path: Path):
    """创建流式章节生成器"""
    import json

    def generator():
        # 解析大纲获取章节
        chapters = _parse_outline_chapters(outline_content)

        if not chapters:
            yield (
                json.dumps(
                    {"type": "error", "message": "无法从大纲中解析出章节"}, ensure_ascii=False
                )
                + "\n"
            )
            return

        # 发送章节列表
        yield (
            json.dumps(
                {
                    "type": "chapters",
                    "data": [{"title": ch["title"], "index": i} for i, ch in enumerate(chapters)],
                },
                ensure_ascii=False,
            )
            + "\n"
        )

        # 检查 LLM 是否可用
        if not deps.rag_service.llm_client.is_available():
            yield (
                json.dumps({"type": "error", "message": "LLM 服务不可用"}, ensure_ascii=False)
                + "\n"
            )
            return

        # 为每个章节生成内容
        for i, chapter in enumerate(chapters):
            chapter_title = chapter["title"]
            chapter_points = (
                "\n".join(f"- {p}" for p in chapter["points"])
                if chapter["points"]
                else "（无具体要点）"
            )

            # 发送开始生成某章节的消息
            yield (
                json.dumps(
                    {"type": "chapter_start", "index": i, "title": chapter_title},
                    ensure_ascii=False,
                )
                + "\n"
            )

            try:
                # 获取 prompt
                system_prompt = get_prompt("project_outline", "chapter_generate_system")
                user_template = get_prompt("project_outline", "chapter_generate_user")

                user_message = (
                    user_template.replace("{project_name}", project_id)
                    .replace("{outline_content}", outline_content)
                    .replace("{chapter_title}", chapter_title)
                    .replace("{chapter_points}", chapter_points)
                )

                # 流式调用 LLM
                response = deps.rag_service.llm_client.client.chat.completions.create(
                    model=deps.rag_service.llm_client.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_message},
                    ],
                    temperature=0.7,
                    stream=True,
                    stream_options={"include_usage": True},
                )

                full_content = f"# {chapter_title}\n\n"
                usage_info = None

                for chunk in response:
                    if hasattr(chunk, "usage") and chunk.usage:
                        usage_info = chunk.usage

                    if chunk.choices and len(chunk.choices) > 0 and chunk.choices[0].delta.content:
                        content = chunk.choices[0].delta.content
                        full_content += content
                        # 发送内容块
                        yield (
                            json.dumps(
                                {"type": "content", "index": i, "chunk": content},
                                ensure_ascii=False,
                            )
                            + "\n"
                        )

                # 生成文件名并保存
                filename = f"{i + 1:02d}_{_sanitize_filename(chapter_title)}.md"
                file_path = project_path / filename
                file_path.write_text(full_content, encoding="utf-8")

                # 发送章节完成消息
                yield (
                    json.dumps(
                        {
                            "type": "chapter_done",
                            "index": i,
                            "title": chapter_title,
                            "filename": filename,
                            "file_id": f"{project_id}/{filename}",
                        },
                        ensure_ascii=False,
                    )
                    + "\n"
                )

                # 记录 token 使用量
                if usage_info:
                    log_token_usage(
                        model=deps.rag_service.llm_client.model,
                        input_tokens=usage_info.prompt_tokens,
                        output_tokens=usage_info.completion_tokens,
                        endpoint="workspace_generate_chapter",
                        user_query=chapter_title,
                        response_type="stream",
                        feature_type="workspace_assistant",
                        additional_info={
                            "project_name": project_id,
                            "chapter_index": i,
                            "chapter_title": chapter_title,
                        },
                    )

            except Exception as e:
                logger.error(f"生成章节 {chapter_title} 失败: {e}")
                yield (
                    json.dumps(
                        {
                            "type": "chapter_error",
                            "index": i,
                            "title": chapter_title,
                            "error": str(e),
                        },
                        ensure_ascii=False,
                    )
                    + "\n"
                )

        # 发送全部完成消息
        yield json.dumps({"type": "done"}, ensure_ascii=False) + "\n"

    return generator


@router.post("/projects/{project_id}/chapters/generate")
async def generate_project_chapters_stream(project_id: str):
    """根据大纲流式生成各章节文件

    Args:
        project_id: 项目 ID
    """
    try:
        workspace_dir = deps.config.workspace_dir
        project_path = Path(workspace_dir) / project_id
        outline_path = project_path / "outline.md"

        # 检查项目是否存在
        if not project_path.exists():

            async def error_gen():
                import json

                yield (
                    json.dumps(
                        {"type": "error", "message": f"项目不存在: {project_id}"},
                        ensure_ascii=False,
                    )
                    + "\n"
                )

            return StreamingResponse(error_gen(), media_type="application/x-ndjson; charset=utf-8")

        # 检查大纲文件是否存在
        if not outline_path.exists():

            async def error_gen():
                import json

                yield (
                    json.dumps(
                        {"type": "error", "message": "大纲文件 outline.md 不存在"},
                        ensure_ascii=False,
                    )
                    + "\n"
                )

            return StreamingResponse(error_gen(), media_type="application/x-ndjson; charset=utf-8")

        # 读取大纲内容
        outline_content = outline_path.read_text(encoding="utf-8")

        # 创建流式生成器
        generator = _create_chapters_stream_generator(project_id, outline_content, project_path)
        headers = {"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}

        return StreamingResponse(
            generator(), media_type="application/x-ndjson; charset=utf-8", headers=headers
        )

    except Exception as e:
        logger.error(f"流式生成章节失败: {e}")
        async def error_gen():
            import json

            yield (
                json.dumps(
                    {"type": "error", "message": f"生成章节失败: {str(e)}"},
                    ensure_ascii=False,
                )
                + "\n"
            )

        return StreamingResponse(error_gen(), media_type="application/x-ndjson; charset=utf-8")


def _create_slides_stream_generator(
    project_id: str, outline_content: str, project_path: Path, custom_requirements: str = ""
):
    """创建流式幻灯片生成器
    
    Args:
        project_id: 项目 ID
        outline_content: 大纲内容
        project_path: 项目路径
        custom_requirements: 用户自定义要求（如颜色、风格等）
    """
    import json

    def generator():
        # 解析大纲获取章节
        chapters = _parse_outline_chapters(outline_content)

        if not chapters:
            yield (
                json.dumps(
                    {"type": "error", "message": "无法从大纲中解析出章节"}, ensure_ascii=False
                )
                + "\n"
            )
            return

        # 为每个章节生成 prompt（章节标题 + 内容）
        # 注意：实际的章节文档内容将在生成时读取并添加到prompt中
        slides_data = []
        for i, chapter in enumerate(chapters):
            chapter_title = chapter["title"]
            chapter_content = chapter.get("content", "").strip()
            chapter_points = chapter.get("points", [])
            
            # 构建基础 prompt：章节标题
            base_prompt = f"生成一个专业的学术演示文稿幻灯片，主题为：{chapter_title}"
            
            # 尝试读取对应的章节.md文件
            safe_title = _sanitize_filename(chapter_title)
            chapter_filename = f"{i + 1:02d}_{safe_title}.md"
            chapter_file_path = project_path / chapter_filename
            
            chapter_doc_content = None
            if chapter_file_path.exists():
                try:
                    chapter_doc_content = chapter_file_path.read_text(encoding="utf-8").strip()
                    logger.info(f"已读取章节文档: {chapter_filename}")
                except Exception as e:
                    logger.warning(f"读取章节文档 {chapter_filename} 失败: {e}")
            
            # 构建完整 prompt
            # 结构：主题 -> 所有要求（集中在前） -> 章节内容
            prompt_parts = [base_prompt]
            
            # 收集所有要求，集中放在前面
            requirements_parts = []
            
            # 1. 默认要求
            requirements_parts.append("要求：简洁、专业、美观，适合学术演示。")
            
            # 2. 自定义要求（如果有）
            if custom_requirements and custom_requirements.strip():
                requirements_parts.append(f"自定义要求：{custom_requirements.strip()}")
            
            # 将所有要求添加到 prompt（在内容之前）
            if requirements_parts:
                prompt_parts.append("\n\n" + "\n\n".join(requirements_parts))
            
            # 3. 章节内容（放在要求之后）
            if chapter_doc_content:
                prompt_parts.append(f"\n\n章节详细内容：\n{chapter_doc_content}")
            # 如果没有章节文档，使用大纲中的内容
            elif chapter_content:
                prompt_parts.append(f"\n\n章节内容：\n{chapter_content}")
            # 如果都没有，使用要点
            elif chapter_points:
                chapter_points_text = "\n".join(f"- {p}" for p in chapter_points)
                prompt_parts.append(f"\n\n要点：\n{chapter_points_text}")
            
            prompt = "".join(prompt_parts)
            
            slides_data.append({
                "title": chapter_title,
                "index": i,
                "prompt": prompt,
            })

        # 发送幻灯片列表
        yield (
            json.dumps(
                {
                    "type": "slides",
                    "data": slides_data,
                },
                ensure_ascii=False,
            )
            + "\n"
        )

        # 检查 Gemini API 是否可用
        try:
            from google import genai
            from google.genai import types
            
            # 从配置读取 Gemini API key
            api_key = deps.config.gemini_api_key
            if not api_key or api_key in ["YOUR_API_KEY_HERE", "YOUR_GEMINI_API_KEY_HERE", "YOUR_LLM_KEY_HERE", "xxx"]:
                yield (
                    json.dumps(
                        {"type": "error", "message": "Gemini API Key 未配置，请在配置文件中设置 gemini.api_key"},
                        ensure_ascii=False,
                    )
                    + "\n"
                )
                return

            # 初始化 Gemini 客户端
            client = genai.Client(api_key=api_key)
            # 使用支持图像生成的模型
            model_name = "gemini-3-pro-image-preview"  # 使用支持图像生成的模型
            
        except ImportError:
            yield (
                json.dumps(
                    {"type": "error", "message": "google-genai 包未安装，请运行: pip install google-genai"},
                    ensure_ascii=False,
                )
                + "\n"
            )
            return
        except Exception as e:
            logger.error(f"初始化 Gemini 客户端失败: {e}")
            yield (
                json.dumps(
                    {"type": "error", "message": f"初始化 Gemini 客户端失败: {str(e)}"},
                    ensure_ascii=False,
                )
                + "\n"
            )
            return

        # 创建 slides 目录
        slides_dir = project_path / "slides"
        slides_dir.mkdir(exist_ok=True)

        # 跟踪已生成的幻灯片文件路径
        generated_slide_paths = []

        # 为每个章节生成幻灯片
        for i, slide_info in enumerate(slides_data):
            chapter_title = slide_info["title"]
            base_prompt = slide_info["prompt"]
            
            # 确定参考窗口大小
            # gemini-3-pro-image-preview: 5 images high fidelity, up to 14 total
            reference_window_size = 5
            logger.info(f"幻灯片 {i + 1}/{len(slides_data)}: 使用模型 {model_name}，参考窗口大小: {reference_window_size}")
            
            # 构建包含参考图片的 contents 数组
            contents = []
            
            # 第一张幻灯片特殊处理：不添加任何参考图片，只使用文本 prompt
            if i == 0:
                # 第一张幻灯片，无参考图片，直接使用文本 prompt
                prompt = base_prompt
                contents.append(prompt)
                logger.info(f"第一张幻灯片，不使用参考图片")
            else:
                # 后续幻灯片：添加参考图片
                start_idx = max(0, i - reference_window_size)
                reference_count = 0
                
                for prev_idx in range(start_idx, i):
                    if prev_idx < len(generated_slide_paths) and generated_slide_paths[prev_idx] is not None:
                        prev_image_path = generated_slide_paths[prev_idx]
                        if prev_image_path.exists():
                            try:
                                with open(prev_image_path, 'rb') as f:
                                    image_bytes = f.read()
                                contents.append(types.Part.from_bytes(
                                    data=image_bytes,
                                    mime_type='image/png'
                                ))
                                reference_count += 1
                                logger.info(f"已加载参考幻灯片 {prev_idx + 1}: {prev_image_path.name}")
                            except Exception as e:
                                logger.warning(f"加载参考幻灯片 {prev_idx + 1} 失败: {e}")
                        else:
                            logger.warning(f"参考幻灯片文件不存在: {prev_image_path}")
                    else:
                        logger.warning(f"参考幻灯片 {prev_idx + 1} 尚未生成，跳过")
                
                # base_prompt 已经包含了：主题 -> 要求（集中在前）-> 章节内容
                # 如果有参考图片，需要在要求部分添加视觉风格要求
                if reference_count > 0:
                    # 在 base_prompt 的要求部分之后添加视觉风格要求
                    # base_prompt 格式：主题\n\n要求...\n\n章节内容...
                    # 需要在"要求"部分之后、"章节内容"之前插入视觉风格要求
                    
                    continuity_instruction = (
                        "视觉风格要求："
                        "\n- 参考提供的之前幻灯片，在整体风格（如配色方案、字体风格、布局结构）上保持适度的一致性"
                        "\n- 但不要过度模仿或完全复制之前的幻灯片，每张幻灯片应该有自己独特的内容呈现方式"
                        "\n- 在保持视觉连贯性的同时，允许适当的创新和变化，确保每张幻灯片都有其独特的视觉表达"
                        "\n- 重点关注内容本身的呈现，而不是机械地复制之前的视觉元素"
                    )
                    
                    # 找到"章节详细内容："或"章节内容："或"要点："的位置，在其之前插入视觉风格要求
                    content_markers = ["章节详细内容：", "章节内容：", "要点："]
                    insert_position = -1
                    for marker in content_markers:
                        if marker in base_prompt:
                            insert_position = base_prompt.find(marker)
                            break
                    
                    if insert_position > 0:
                        # 在章节内容之前插入视觉风格要求
                        prompt = (
                            base_prompt[:insert_position].rstrip() +
                            "\n\n" + continuity_instruction +
                            base_prompt[insert_position:]
                        )
                    else:
                        # 如果找不到章节内容标记，直接追加到末尾（不应该发生）
                        prompt = base_prompt + "\n\n" + continuity_instruction
                    
                    logger.info(f"幻灯片 {i + 1} 使用 {reference_count} 张参考图片")
                else:
                    prompt = base_prompt
                
                # 将文本 prompt 添加到 contents（在参考图片之后）
                contents.append(prompt)

            # 发送开始生成某幻灯片的消息
            yield (
                json.dumps(
                    {"type": "slide_start", "index": i, "title": chapter_title, "prompt": prompt},
                    ensure_ascii=False,
                )
                + "\n"
            )

            try:
                # 发送进度消息
                # 第一张幻灯片：contents 只有 prompt（字符串），所以 len(contents) = 1
                # 后续幻灯片：contents 包含参考图片（Part对象）+ prompt（字符串），所以 len(contents) > 1
                if i == 0:
                    reference_info = "（第一张幻灯片，无参考图片）"
                else:
                    # 计算参考图片数量（contents 中除了最后一个 prompt，其他都是图片）
                    ref_count = len(contents) - 1 if len(contents) > 1 else 0
                    reference_info = f"（使用 {ref_count} 张参考图片）" if ref_count > 0 else ""
                yield (
                    json.dumps(
                        {"type": "slide_progress", "index": i, "message": f"正在调用 Gemini API 生成图片{reference_info}..."},
                        ensure_ascii=False,
                    )
                    + "\n"
                )

                # 配置图片生成参数
                image_config = types.ImageConfig(
                    aspectRatio="16:9",  # 幻灯片使用 16:9 比例
                )

                # 重试机制：最多重试3次，指数退避
                import time
                max_retries = 3
                retry_delay = 1  # 初始延迟1秒
                response = None
                last_error = None
                
                for attempt in range(max_retries):
                    try:
                        # 第一张幻灯片特殊处理：确保 contents 格式正确
                        if i == 0:
                            # 第一张幻灯片：contents 应该只包含 prompt 文本（字符串）
                            if len(contents) != 1 or not isinstance(contents[0], str):
                                raise Exception(f"第一张幻灯片 contents 格式错误: 期望长度为1的字符串列表，实际: {len(contents)} 个元素")
                            logger.info(f"第一张幻灯片 contents: 仅包含文本 prompt，长度: {len(contents[0])}")
                        else:
                            # 后续幻灯片：contents 应该包含参考图片（Part对象）+ prompt（字符串）
                            if len(contents) < 1:
                                raise Exception(f"后续幻灯片 contents 为空")
                            # 最后一个应该是 prompt（字符串），前面的应该是图片（Part对象）
                            if not isinstance(contents[-1], str):
                                raise Exception(f"后续幻灯片 contents 最后一个元素应该是字符串（prompt），实际类型: {type(contents[-1])}")
                            logger.info(f"幻灯片 {i + 1} contents: {len(contents) - 1} 张参考图片 + 1 个文本 prompt")
                        
                        # 调用 Gemini API 生成图片
                        response = client.models.generate_content(
                            model=model_name,
                            contents=contents,  # 第一张：仅文本；后续：图片+文本
                            config=types.GenerateContentConfig(
                                response_modalities=["IMAGE"],
                                image_config=image_config,
                            ),
                        )
                        
                        # 检查响应是否有效
                        if response is None:
                            raise Exception("API返回空响应")
                        
                        # 检查 response.parts 是否存在且可迭代
                        if not hasattr(response, 'parts') or response.parts is None:
                            raise Exception("响应中缺少parts属性或parts为None")
                        
                        # 成功获取响应，跳出重试循环
                        break
                        
                    except Exception as api_error:
                        last_error = api_error
                        error_str = str(api_error).lower()
                        
                        # 判断是否可重试的错误
                        is_retryable = (
                            "network" in error_str or
                            "timeout" in error_str or
                            "connection" in error_str or
                            "service unavailable" in error_str or
                            "rate limit" in error_str or
                            "429" in error_str or
                            "500" in error_str or
                            "502" in error_str or
                            "503" in error_str or
                            "504" in error_str or
                            "none" in error_str or
                            "null" in error_str
                        )
                        
                        if not is_retryable or attempt == max_retries - 1:
                            # 不可重试或已达到最大重试次数，抛出异常
                            raise api_error
                        
                        # 可重试，等待后重试
                        logger.warning(f"生成幻灯片 {i + 1} 第 {attempt + 1} 次尝试失败: {api_error}，{retry_delay}秒后重试...")
                        time.sleep(retry_delay)
                        retry_delay *= 2  # 指数退避：1s, 2s, 4s
                
                # 最终检查响应有效性
                if response is None:
                    raise Exception(f"API调用失败，已重试{max_retries}次: {last_error}")
                
                if not hasattr(response, 'parts') or response.parts is None:
                    raise Exception("响应中缺少parts属性或parts为None，无法提取图片")

                # 发送进度消息
                yield (
                    json.dumps(
                        {"type": "slide_progress", "index": i, "message": "正在保存图片..."},
                        ensure_ascii=False,
                    )
                    + "\n"
                )

                # 从响应中提取图片并保存
                image_saved = False
                saved_file_path = None
                
                # 安全地迭代 response.parts
                parts_to_iterate = []
                try:
                    if response.parts is not None:
                        # 尝试转换为列表以确保可迭代
                        if isinstance(response.parts, (list, tuple)):
                            parts_to_iterate = list(response.parts)
                        else:
                            # 尝试迭代
                            parts_to_iterate = list(response.parts)
                except (TypeError, AttributeError) as e:
                    logger.error(f"无法迭代响应parts: {e}, response.parts类型: {type(response.parts)}")
                    raise Exception(f"响应parts格式错误，无法迭代: {e}")
                
                for part in parts_to_iterate:
                    if part is None:
                        continue
                    if hasattr(part, 'inline_data') and part.inline_data is not None:
                        try:
                            # 使用 part.as_image() 方法获取图片对象
                            image = part.as_image()
                            
                            # 生成文件名
                            safe_title = _sanitize_filename(chapter_title)
                            filename = f"slide_{i + 1:02d}_{safe_title}.png"
                            file_path = slides_dir / filename

                            # 使用图片对象的 save() 方法保存
                            image.save(str(file_path))
                            image_saved = True
                            saved_file_path = file_path

                            # 构建访问URL
                            image_url = f"/api/workspace/slides/{project_id}/{filename}"

                            # 发送幻灯片完成消息
                            yield (
                                json.dumps(
                                    {
                                        "type": "slide_done",
                                        "index": i,
                                        "title": chapter_title,
                                        "filename": filename,
                                        "url": image_url,
                                    },
                                    ensure_ascii=False,
                                )
                                + "\n"
                            )
                            break
                        except Exception as e:
                            logger.error(f"保存幻灯片图片失败: {e}")
                            continue

                if not image_saved:
                    # 如果没有找到图片数据，尝试检查文本响应
                    error_msg = "未能从响应中提取图片数据"
                    try:
                        if response and hasattr(response, 'parts') and response.parts is not None:
                            # 安全地迭代 response.parts
                            parts_list = []
                            try:
                                if isinstance(response.parts, (list, tuple)):
                                    parts_list = list(response.parts)
                                else:
                                    parts_list = list(response.parts)
                            except (TypeError, AttributeError):
                                parts_list = []
                            
                            text_parts = [part.text for part in parts_list if hasattr(part, 'text') and part.text is not None]
                            if text_parts:
                                error_msg += f"。响应文本: {''.join(text_parts)}"
                    except Exception as e:
                        logger.warning(f"提取响应文本失败: {e}")
                    # 抛出异常，由外层异常处理器统一处理 generated_slide_paths
                    raise Exception(error_msg)
                else:
                    # 成功生成，记录文件路径供后续幻灯片参考
                    generated_slide_paths.append(saved_file_path)
                    logger.info(f"幻灯片 {i + 1} 已保存: {saved_file_path.name}")

            except Exception as e:
                logger.error(f"生成幻灯片 {chapter_title} 失败: {e}")
                # 记录失败，但不中断后续幻灯片的生成
                # 如果 generated_slide_paths 长度不够，补充 None
                while len(generated_slide_paths) <= i:
                    generated_slide_paths.append(None)
                generated_slide_paths[i] = None
                
                yield (
                    json.dumps(
                        {
                            "type": "slide_error",
                            "index": i,
                            "title": chapter_title,
                            "error": str(e),
                        },
                        ensure_ascii=False,
                    )
                    + "\n"
                )

        # 发送全部完成消息
        yield json.dumps({"type": "done"}, ensure_ascii=False) + "\n"

    return generator


@router.get("/projects/{project_id}/slides/preview", response_model=PreviewSlidesPromptsResponse)
async def preview_slides_prompts(
    project_id: str,
    custom_requirements: str = Query("", description="用户自定义要求（如颜色、风格等），用于预览"),
):
    """预览幻灯片生成所需的 prompts 列表

    Args:
        project_id: 项目 ID
        custom_requirements: 用户自定义要求（如颜色、风格等），用于预览

    Returns:
        包含 prompts 列表的响应
    """
    try:
        workspace_dir = deps.config.workspace_dir
        project_path = Path(workspace_dir) / project_id
        outline_path = project_path / "outline.md"

        # 检查项目是否存在
        if not project_path.exists():
            return PreviewSlidesPromptsResponse(
                success=False,
                prompts=[],
                total=0,
                error=f"项目不存在: {project_id}",
            )

        # 检查大纲文件是否存在
        if not outline_path.exists():
            return PreviewSlidesPromptsResponse(
                success=False,
                prompts=[],
                total=0,
                error="大纲文件 outline.md 不存在",
            )

        # 读取大纲内容
        outline_content = outline_path.read_text(encoding="utf-8")

        # 解析大纲获取章节
        chapters = _parse_outline_chapters(outline_content)

        if not chapters:
            return PreviewSlidesPromptsResponse(
                success=False,
                prompts=[],
                total=0,
                error="无法从大纲中解析出章节",
            )

        # 检查是否已生成章节文件
        chapter_files_exist = False
        for i, chapter in enumerate(chapters):
            safe_title = _sanitize_filename(chapter["title"])
            chapter_filename = f"{i + 1:02d}_{safe_title}.md"
            chapter_file_path = project_path / chapter_filename
            if chapter_file_path.exists():
                chapter_files_exist = True
                break

        if not chapter_files_exist:
            return PreviewSlidesPromptsResponse(
                success=False,
                prompts=[],
                total=0,
                error="请先生成章节文件，然后再生成幻灯片。章节文件格式：01_章节名.md",
            )

        # 为每个章节生成 prompt（章节标题 + 内容）
        from lifetrace.schemas.workspace import SlidePromptPreview

        prompts = []
        for i, chapter in enumerate(chapters):
            chapter_title = chapter["title"]
            chapter_content = chapter.get("content", "").strip()
            chapter_points = chapter.get("points", [])
            
            # 构建基础 prompt：章节标题
            base_prompt = f"生成一个专业的学术演示文稿幻灯片，主题为：{chapter_title}"
            
            # 尝试读取对应的章节.md文件
            safe_title = _sanitize_filename(chapter_title)
            chapter_filename = f"{i + 1:02d}_{safe_title}.md"
            chapter_file_path = project_path / chapter_filename
            
            chapter_doc_content = None
            if chapter_file_path.exists():
                try:
                    chapter_doc_content = chapter_file_path.read_text(encoding="utf-8").strip()
                except Exception as e:
                    logger.warning(f"读取章节文档 {chapter_filename} 失败: {e}")
            
            # 构建完整 prompt
            # 结构：主题 -> 所有要求（集中在前） -> 章节内容
            prompt_parts = [base_prompt]
            
            # 收集所有要求，集中放在前面
            requirements_parts = []
            
            # 1. 默认要求
            requirements_parts.append("要求：简洁、专业、美观，适合学术演示。")
            
            # 2. 自定义要求（如果有，预览接口也支持）
            if custom_requirements and custom_requirements.strip():
                requirements_parts.append(f"自定义要求：{custom_requirements.strip()}")
            
            # 将所有要求添加到 prompt（在内容之前）
            if requirements_parts:
                prompt_parts.append("\n\n" + "\n\n".join(requirements_parts))
            
            # 章节内容（放在要求之后）
            if chapter_doc_content:
                prompt_parts.append(f"\n\n章节详细内容：\n{chapter_doc_content}")
            # 如果没有章节文档，使用大纲中的内容
            elif chapter_content:
                prompt_parts.append(f"\n\n章节内容：\n{chapter_content}")
            # 如果都没有，使用要点
            elif chapter_points:
                chapter_points_text = "\n".join(f"- {p}" for p in chapter_points)
                prompt_parts.append(f"\n\n要点：\n{chapter_points_text}")
            
            prompt = "".join(prompt_parts)
            
            prompts.append(
                SlidePromptPreview(
                    title=chapter_title,
                    index=i,
                    prompt=prompt,
                )
            )

        return PreviewSlidesPromptsResponse(
            success=True,
            prompts=prompts,
            total=len(prompts),
        )

    except Exception as e:
        logger.error(f"预览幻灯片 prompts 失败: {e}")
        return PreviewSlidesPromptsResponse(
            success=False,
            prompts=[],
            total=0,
            error=str(e),
        )


@router.post("/projects/{project_id}/slides/generate")
async def generate_project_slides_stream(
    project_id: str,
    custom_requirements: str = Query("", description="用户自定义要求（如颜色、风格等）"),
):
    """根据大纲流式生成幻灯片图片

    Args:
        project_id: 项目 ID
        custom_requirements: 用户自定义要求（如颜色、风格等）
    """
    try:
        workspace_dir = deps.config.workspace_dir
        project_path = Path(workspace_dir) / project_id
        outline_path = project_path / "outline.md"

        # 检查项目是否存在
        if not project_path.exists():
            async def error_gen():
                import json

                yield (
                    json.dumps(
                        {"type": "error", "message": f"项目不存在: {project_id}"},
                        ensure_ascii=False,
                    )
                    + "\n"
                )

            return StreamingResponse(error_gen(), media_type="application/x-ndjson; charset=utf-8")

        # 检查大纲文件是否存在
        if not outline_path.exists():
            async def error_gen():
                import json

                yield (
                    json.dumps(
                        {"type": "error", "message": "大纲文件 outline.md 不存在"},
                        ensure_ascii=False,
                    )
                    + "\n"
                )

            return StreamingResponse(error_gen(), media_type="application/x-ndjson; charset=utf-8")

        # 读取大纲内容
        outline_content = outline_path.read_text(encoding="utf-8")
        
        # 解析大纲获取章节列表
        chapters = _parse_outline_chapters(outline_content)
        
        # 检查是否已生成章节文件（前置要求）
        chapter_files_exist = False
        for i, chapter in enumerate(chapters):
            safe_title = _sanitize_filename(chapter["title"])
            chapter_filename = f"{i + 1:02d}_{safe_title}.md"
            chapter_file_path = project_path / chapter_filename
            if chapter_file_path.exists():
                chapter_files_exist = True
                break
        
        if not chapter_files_exist:
            async def error_gen():
                import json

                yield (
                    json.dumps(
                        {"type": "error", "message": "请先生成章节文件，然后再生成幻灯片。章节文件格式：01_章节名.md"},
                        ensure_ascii=False,
                    )
                    + "\n"
                )

            return StreamingResponse(error_gen(), media_type="application/x-ndjson; charset=utf-8")

        # 创建流式生成器，传递自定义要求
        generator = _create_slides_stream_generator(
            project_id, outline_content, project_path, custom_requirements
        )
        headers = {"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}

        return StreamingResponse(
            generator(), media_type="application/x-ndjson; charset=utf-8", headers=headers
        )

    except Exception as e:
        logger.error(f"流式生成幻灯片失败: {e}")
        async def error_gen():
            import json

            yield (
                json.dumps(
                    {"type": "error", "message": f"生成幻灯片失败: {str(e)}"},
                    ensure_ascii=False,
                )
                + "\n"
            )

        return StreamingResponse(error_gen(), media_type="application/x-ndjson; charset=utf-8")
        raise HTTPException(status_code=500, detail=str(e)) from e


@router.post("/projects", response_model=CreateWorkspaceProjectResponse)
async def create_workspace_project(request: CreateWorkspaceProjectRequest):
    """创建工作区项目（一级文件夹）"""
    try:
        if not request.name or not request.name.strip():
            return CreateWorkspaceProjectResponse(success=False, error="项目名称不能为空")

        project_name = request.name.strip()

        # 检查项目名称是否包含非法字符
        invalid_chars = ["/", "\\", ":", "*", "?", '"', "<", ">", "|"]
        if any(char in project_name for char in invalid_chars):
            return CreateWorkspaceProjectResponse(
                success=False, error=f"项目名称包含非法字符: {invalid_chars}"
            )

        workspace_dir = deps.config.workspace_dir
        project_path = Path(workspace_dir) / project_name

        # 检查项目是否已存在
        if project_path.exists():
            return CreateWorkspaceProjectResponse(
                success=False, error=f"项目已存在: {project_name}"
            )

        # 创建项目文件夹
        project_path.mkdir(parents=True, exist_ok=True)

        # 复制项目模板文件夹结构
        template_dir = Path(__file__).parent.parent / "templates" / "paper"
        if template_dir.exists():
            for item in template_dir.iterdir():
                if item.is_dir():
                    dest = project_path / item.name
                    shutil.copytree(item, dest)
                    logger.debug(f"复制模板文件夹: {item.name} -> {dest}")

        # 先生成模板版本的大纲文件（不调用 LLM）
        outline_content = _generate_outline_template(project_name, request.project_type)
        outline_path = project_path / "outline.md"
        outline_path.write_text(outline_content, encoding="utf-8")

        logger.info(f"创建工作区项目成功: {project_name}, 类型: {request.project_type.value}")

        return CreateWorkspaceProjectResponse(
            success=True,
            project_id=project_name,
            project_name=project_name,
        )

    except Exception as e:
        logger.error(f"创建工作区项目失败: {e}")
        return CreateWorkspaceProjectResponse(success=False, error=str(e))


@router.post("/projects/rename", response_model=RenameWorkspaceProjectResponse)
async def rename_workspace_project(request: RenameWorkspaceProjectRequest):
    """重命名工作区项目"""
    try:
        if not request.new_name or not request.new_name.strip():
            return RenameWorkspaceProjectResponse(
                success=False,
                old_id=request.project_id,
                new_id=request.project_id,
                new_name=request.new_name,
                error="新项目名称不能为空",
            )

        new_name = request.new_name.strip()

        # 检查新项目名称是否包含非法字符
        invalid_chars = ["/", "\\", ":", "*", "?", '"', "<", ">", "|"]
        if any(char in new_name for char in invalid_chars):
            return RenameWorkspaceProjectResponse(
                success=False,
                old_id=request.project_id,
                new_id=request.project_id,
                new_name=new_name,
                error=f"项目名称包含非法字符: {invalid_chars}",
            )

        workspace_dir = deps.config.workspace_dir
        old_path = Path(workspace_dir) / request.project_id
        new_path = Path(workspace_dir) / new_name

        # 检查旧项目是否存在
        if not old_path.exists():
            return RenameWorkspaceProjectResponse(
                success=False,
                old_id=request.project_id,
                new_id=request.project_id,
                new_name=new_name,
                error=f"项目不存在: {request.project_id}",
            )

        # 检查新项目名是否已存在
        if new_path.exists() and new_path != old_path:
            return RenameWorkspaceProjectResponse(
                success=False,
                old_id=request.project_id,
                new_id=request.project_id,
                new_name=new_name,
                error=f"项目已存在: {new_name}",
            )

        # 重命名
        old_path.rename(new_path)

        logger.info(f"重命名工作区项目成功: {request.project_id} -> {new_name}")

        return RenameWorkspaceProjectResponse(
            success=True,
            old_id=request.project_id,
            new_id=new_name,
            new_name=new_name,
        )

    except Exception as e:
        logger.error(f"重命名工作区项目失败: {e}")
        return RenameWorkspaceProjectResponse(
            success=False,
            old_id=request.project_id,
            new_id=request.project_id,
            new_name=request.new_name,
            error=str(e),
        )


@router.post("/projects/delete", response_model=DeleteWorkspaceProjectResponse)
async def delete_workspace_project(request: DeleteWorkspaceProjectRequest):
    """删除工作区项目（一级文件夹及其所有内容）"""
    try:
        if not request.project_id:
            return DeleteWorkspaceProjectResponse(
                success=False, project_id=request.project_id, error="项目 ID 不能为空"
            )

        workspace_dir = deps.config.workspace_dir
        project_path = Path(workspace_dir) / request.project_id

        # 安全检查：确保路径在 workspace 目录内
        try:
            project_path.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            return DeleteWorkspaceProjectResponse(
                success=False, project_id=request.project_id, error="禁止访问工作区外的目录"
            )

        # 检查项目是否存在
        if not project_path.exists():
            return DeleteWorkspaceProjectResponse(
                success=False,
                project_id=request.project_id,
                error=f"项目不存在: {request.project_id}",
            )

        # 检查是否为目录
        if not project_path.is_dir():
            return DeleteWorkspaceProjectResponse(
                success=False,
                project_id=request.project_id,
                error=f"不是有效的项目: {request.project_id}",
            )

        # 递归删除项目文件夹
        shutil.rmtree(project_path)

        logger.info(f"删除工作区项目成功: {request.project_id}")

        return DeleteWorkspaceProjectResponse(success=True, project_id=request.project_id)

    except Exception as e:
        logger.error(f"删除工作区项目失败: {e}")
        return DeleteWorkspaceProjectResponse(
            success=False, project_id=request.project_id, error=str(e)
        )


@router.get("/projects/{project_id}/files", response_model=WorkspaceFilesResponse)
async def get_project_files(project_id: str):
    """获取指定项目下的文件列表"""
    try:
        workspace_dir = deps.config.workspace_dir
        project_path = Path(workspace_dir) / project_id

        # 安全检查：确保路径在 workspace 目录内
        try:
            project_path.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            raise HTTPException(status_code=403, detail="禁止访问工作区外的目录") from None

        # 检查项目目录是否存在
        if not project_path.exists():
            logger.warning(f"项目目录不存在: {project_path}")
            return WorkspaceFilesResponse(files=[], total=0)

        # 构建文件树（在项目目录下）
        files = build_file_tree_for_project(str(project_path), project_id)
        total = count_files(files)

        logger.debug(f"获取项目 {project_id} 文件列表成功，共 {total} 个文件")
        return WorkspaceFilesResponse(files=files, total=total)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取项目文件列表失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


def _file_sort_key(item: Path, is_root: bool = False) -> tuple:
    """文件排序键：outline.md 在根目录置顶，然后文件夹优先，最后按名称排序"""
    is_outline = is_root and item.name.lower() == "outline.md"
    return (not is_outline, not item.is_dir(), item.name.lower())


def build_file_tree_for_project(
    root_path: str, project_id: str, parent_id: str = "", is_root: bool = True
) -> list[FileNode]:
    """递归构建项目内的文件树

    Args:
        root_path: 根目录路径
        project_id: 项目 ID
        parent_id: 父节点ID
        is_root: 是否为项目根目录

    Returns:
        文件节点列表
    """
    nodes = []
    root = Path(root_path)

    if not root.exists():
        return nodes

    # 获取所有子项并排序（outline.md 置顶，然后文件夹优先，最后按名称排序）
    items = sorted(root.iterdir(), key=lambda x: _file_sort_key(x, is_root))

    workspace_dir = deps.config.workspace_dir
    project_path = Path(workspace_dir) / project_id

    for item in items:
        # 跳过隐藏文件和目录
        if item.name.startswith("."):
            continue

        # 生成节点ID（相对于 workspace 的路径，包含项目 ID）
        relative_path = item.relative_to(project_path)
        node_id = f"{project_id}/{str(relative_path).replace(os.sep, '/')}"

        # 检查是否为受保护的文件（outline.md 在根目录）
        is_protected = is_root and item.name.lower() == "outline.md"

        if item.is_dir():
            # 递归处理子目录（不再是根目录）
            children = build_file_tree_for_project(str(item), project_id, node_id, is_root=False)
            node = FileNode(
                id=node_id,
                name=item.name,
                type="folder",
                children=children,
                parent_id=parent_id if parent_id else None,
            )
        else:
            node = FileNode(
                id=node_id,
                name=item.name,
                type="file",
                parent_id=parent_id if parent_id else None,
                is_protected=is_protected,
            )

        nodes.append(node)

    return nodes


# ==================== 原有文件管理 API ====================


def build_file_tree(root_path: str, parent_id: str = "") -> list[FileNode]:
    """递归构建文件树

    Args:
        root_path: 根目录路径
        parent_id: 父节点ID

    Returns:
        文件节点列表
    """
    nodes = []
    root = Path(root_path)

    if not root.exists():
        return nodes

    # 获取所有子项并排序（文件夹优先，然后按名称排序）
    items = sorted(root.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))

    for item in items:
        # 跳过隐藏文件和目录
        if item.name.startswith("."):
            continue

        # 生成节点ID（使用相对路径）
        relative_path = item.relative_to(Path(deps.config.workspace_dir))
        node_id = str(relative_path).replace(os.sep, "/")

        if item.is_dir():
            # 递归处理子目录
            children = build_file_tree(str(item), node_id)
            node = FileNode(
                id=node_id,
                name=item.name,
                type="folder",
                children=children,
                parent_id=parent_id if parent_id else None,
            )
        else:
            node = FileNode(
                id=node_id,
                name=item.name,
                type="file",
                parent_id=parent_id if parent_id else None,
            )

        nodes.append(node)

    return nodes


def count_files(nodes: list[FileNode]) -> int:
    """统计文件总数（不包括文件夹）"""
    count = 0
    for node in nodes:
        if node.type == "file":
            count += 1
        elif node.children:
            count += count_files(node.children)
    return count


def extract_text_from_docx(file_path: Path) -> str:
    """从 docx 文件中提取文本内容

    Args:
        file_path: docx 文件路径

    Returns:
        提取的文本内容
    """
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except ImportError:
        logger.error("python-docx 未安装，无法解析 docx 文件")
        raise
    except Exception as e:
        logger.error(f"解析 docx 文件失败: {e}")
        raise


def extract_text_from_doc(file_path: Path) -> str:
    """从 doc 文件中提取文本内容（使用 antiword 或转换方法）

    Args:
        file_path: doc 文件路径

    Returns:
        提取的文本内容

    Note:
        老版本 .doc 文件较难处理，这里尝试用 python-docx 读取，
        如果失败则返回提示信息
    """
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception:
        # 老版本 .doc 文件 python-docx 无法直接读取
        return f"[无法解析旧版 .doc 文件: {file_path.name}，请转换为 .docx 格式]"


ALLOWED_EXTENSIONS = {".txt", ".md", ".doc", ".docx"}


def _get_unique_file_path(target_dir: Path, filename: str) -> Path:
    """获取唯一的文件路径，如果文件已存在则添加数字后缀"""
    target_file = target_dir / filename
    if not target_file.exists():
        return target_file

    base_name = target_file.stem
    ext = target_file.suffix
    counter = 1
    while target_file.exists():
        target_file = target_dir / f"{base_name}_{counter}{ext}"
        counter += 1
    return target_file


def _decode_text_content(file_content: bytes) -> str:
    """尝试多种编码解码文本内容"""
    encodings = ["utf-8", "gbk", "latin-1"]
    for encoding in encodings:
        try:
            return file_content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_content.decode("latin-1")


def _process_text_file(target_file: Path, file_content: bytes) -> str:
    """处理文本文件（txt, md）"""
    content = _decode_text_content(file_content)
    target_file.write_text(content, encoding="utf-8")
    return content


def _process_doc_file(target_file: Path, file_content: bytes, is_docx: bool) -> tuple[Path, str]:
    """处理 Word 文档文件（doc, docx），返回 (最终文件路径, 内容)"""
    target_file.write_bytes(file_content)
    content = extract_text_from_docx(target_file) if is_docx else extract_text_from_doc(target_file)
    md_file = target_file.with_suffix(".md")
    md_file.write_text(content, encoding="utf-8")
    target_file.unlink()
    return md_file, content


def _validate_upload_target(folder: str, workspace_dir: str) -> tuple[Path | None, str | None]:
    """验证并构建上传目标目录，返回 (目标目录, 错误信息)"""
    target_dir = Path(workspace_dir) / folder if folder else Path(workspace_dir)
    target_dir.mkdir(parents=True, exist_ok=True)

    try:
        target_dir.resolve().relative_to(Path(workspace_dir).resolve())
    except ValueError:
        return None, "禁止访问工作区外的目录"

    return target_dir, None


@router.post("/upload", response_model=UploadFileResponse)
async def upload_file(
    file: UploadFile = File(...),
    folder: str = Query("", description="目标文件夹路径（相对于 workspace）"),
):
    """上传文件到工作区

    支持的文件格式：txt, md, doc, docx

    Args:
        file: 上传的文件
        folder: 目标文件夹路径（可选）
    """
    try:
        if not file.filename:
            return UploadFileResponse(success=False, error="文件名不能为空")

        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            return UploadFileResponse(
                success=False,
                error=f"不支持的文件格式: {file_ext}，支持的格式: {', '.join(ALLOWED_EXTENSIONS)}",
            )

        workspace_dir = deps.config.workspace_dir
        target_dir, error = _validate_upload_target(folder, workspace_dir)
        if error:
            return UploadFileResponse(success=False, error=error)

        target_file = _get_unique_file_path(target_dir, file.filename)
        file_content = await file.read()

        # 根据文件类型处理内容
        if file_ext in {".txt", ".md"}:
            content = _process_text_file(target_file, file_content)
        elif file_ext in {".docx", ".doc"}:
            target_file, content = _process_doc_file(target_file, file_content, file_ext == ".docx")
        else:
            content = ""

        relative_path = target_file.relative_to(Path(workspace_dir))
        file_id = str(relative_path).replace(os.sep, "/")
        logger.info(f"文件上传成功: {file_id}")

        return UploadFileResponse(
            success=True,
            file_id=file_id,
            file_name=target_file.name,
            content=content,
        )

    except Exception as e:
        logger.error(f"文件上传失败: {e}")
        return UploadFileResponse(success=False, error=str(e))


ALLOWED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg"}


@router.post("/upload-image", response_model=UploadImageResponse)
async def upload_image(
    file: UploadFile = File(...),
    project_id: str = Query(..., description="项目ID"),
):
    """上传图片到工作区项目的 images 文件夹

    Args:
        file: 上传的图片文件
        project_id: 项目ID

    Returns:
        上传结果，包含图片访问URL
    """
    try:
        if not file.filename:
            return UploadImageResponse(success=False, error="文件名不能为空")

        # 验证文件扩展名
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_IMAGE_EXTENSIONS:
            return UploadImageResponse(
                success=False,
                error=f"不支持的图片格式: {file_ext}，支持的格式: {', '.join(ALLOWED_IMAGE_EXTENSIONS)}",
            )

        workspace_dir = deps.config.workspace_dir
        project_path = Path(workspace_dir) / project_id

        # 验证项目目录存在
        if not project_path.exists():
            return UploadImageResponse(
                success=False,
                error=f"项目不存在: {project_id}",
            )

        # 创建 images 目录
        images_dir = project_path / "images"
        images_dir.mkdir(exist_ok=True)

        # 生成唯一文件名（添加时间戳避免冲突）
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = Path(file.filename).stem[:50]  # 限制文件名长度
        unique_filename = f"{safe_filename}_{timestamp}{file_ext}"
        target_file = images_dir / unique_filename

        # 读取并保存文件
        file_content = await file.read()

        # 验证文件大小（使用已有的 MAX_FILE_SIZE 常量）
        from lifetrace.util.logging_config import get_logger as gl
        MAX_IMAGE_SIZE = 10 * 1024 * 1024  # 10MB for images
        if len(file_content) > MAX_IMAGE_SIZE:
            return UploadImageResponse(
                success=False,
                error=f"图片文件过大，最大支持 {MAX_IMAGE_SIZE / 1024 / 1024}MB",
            )

        target_file.write_bytes(file_content)

        # 构建访问URL（相对于API的路径）
        image_url = f"/api/workspace/images/{project_id}/{unique_filename}"

        logger.info(f"图片上传成功: {project_id}/images/{unique_filename}")

        return UploadImageResponse(
            success=True,
            url=image_url,
            filename=unique_filename,
        )

    except Exception as e:
        logger.error(f"图片上传失败: {e}")
        return UploadImageResponse(success=False, error=str(e))


@router.get("/images/{project_id}/{filename}")
async def get_image(project_id: str, filename: str):
    """获取项目中的图片文件

    Args:
        project_id: 项目ID
        filename: 图片文件名

    Returns:
        图片文件流
    """
    try:
        workspace_dir = deps.config.workspace_dir
        image_path = Path(workspace_dir) / project_id / "images" / filename

        # 安全检查：确保路径在 workspace 目录内
        try:
            image_path.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            raise HTTPException(status_code=403, detail="禁止访问工作区外的文件") from None

        # 检查文件是否存在
        if not image_path.exists() or not image_path.is_file():
            raise HTTPException(status_code=404, detail="图片不存在")

        # 验证是否为图片文件
        file_ext = image_path.suffix.lower()
        if file_ext not in ALLOWED_IMAGE_EXTENSIONS:
            raise HTTPException(status_code=400, detail="不是有效的图片文件")

        # 确定 MIME 类型
        mime_types = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".gif": "image/gif",
            ".webp": "image/webp",
            ".svg": "image/svg+xml",
        }
        media_type = mime_types.get(file_ext, "application/octet-stream")

        # 创建文件流响应
        def iter_file():
            with open(image_path, "rb") as f:
                yield from f

        return StreamingResponse(iter_file(), media_type=media_type)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取图片失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


@router.get("/slides/{project_id}/images", response_model=SlidesImagesResponse)
async def get_slides_images(project_id: str):
    """获取项目中 slides 文件夹下的所有图片

    Args:
        project_id: 项目ID

    Returns:
        图片列表响应
    """
    try:
        workspace_dir = deps.config.workspace_dir
        slides_dir = Path(workspace_dir) / project_id / "slides"

        # 安全检查：确保路径在 workspace 目录内
        try:
            slides_dir.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            return SlidesImagesResponse(
                success=False,
                images=[],
                error="禁止访问工作区外的目录",
            )

        # 检查 slides 目录是否存在
        if not slides_dir.exists() or not slides_dir.is_dir():
            return SlidesImagesResponse(
                success=True,
                images=[],
            )

        # 获取所有图片文件
        image_files = []
        for file_path in sorted(slides_dir.iterdir()):
            if file_path.is_file():
                file_ext = file_path.suffix.lower()
                if file_ext in ALLOWED_IMAGE_EXTENSIONS:
                    # 构建访问URL
                    image_url = f"/api/workspace/slides/{project_id}/{file_path.name}"
                    image_files.append(
                        SlideImageInfo(
                            url=image_url,
                            name=file_path.name,
                        )
                    )

        logger.info(f"获取项目 {project_id} 的 slides 图片列表成功，共 {len(image_files)} 张")
        return SlidesImagesResponse(
            success=True,
            images=image_files,
        )

    except Exception as e:
        logger.error(f"获取 slides 图片列表失败: {e}")
        return SlidesImagesResponse(
            success=False,
            images=[],
            error=str(e),
        )


@router.get("/slides/{project_id}/{filename}")
async def get_slide_image(project_id: str, filename: str):
    """获取项目中 slides 文件夹下的图片文件

    Args:
        project_id: 项目ID
        filename: 图片文件名

    Returns:
        图片文件流
    """
    try:
        workspace_dir = deps.config.workspace_dir
        image_path = Path(workspace_dir) / project_id / "slides" / filename

        # 安全检查：确保路径在 workspace 目录内
        try:
            image_path.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            raise HTTPException(status_code=403, detail="禁止访问工作区外的文件") from None

        # 检查文件是否存在
        if not image_path.exists() or not image_path.is_file():
            raise HTTPException(status_code=404, detail="图片不存在")

        # 验证是否为图片文件
        file_ext = image_path.suffix.lower()
        if file_ext not in ALLOWED_IMAGE_EXTENSIONS:
            raise HTTPException(status_code=400, detail="不是有效的图片文件")

        # 确定 MIME 类型
        mime_types = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".gif": "image/gif",
            ".webp": "image/webp",
            ".svg": "image/svg+xml",
        }
        media_type = mime_types.get(file_ext, "application/octet-stream")

        # 创建文件流响应
        def iter_file():
            with open(image_path, "rb") as f:
                yield from f

        return StreamingResponse(iter_file(), media_type=media_type)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取 slides 图片失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


@router.post("/create", response_model=CreateFileResponse)
async def create_file(request: CreateFileRequest):
    """创建新文件

    Args:
        request: 创建文件请求，包含文件名和可选的文件夹路径
    """
    try:
        if not request.file_name:
            return CreateFileResponse(success=False, error="文件名不能为空")

        workspace_dir = deps.config.workspace_dir

        # 构建目标文件路径
        if request.folder:
            target_dir = Path(workspace_dir) / request.folder
        else:
            target_dir = Path(workspace_dir)

        # 确保目标目录存在
        target_dir.mkdir(parents=True, exist_ok=True)

        # 安全检查：确保目标目录在 workspace 内
        try:
            target_dir.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            return CreateFileResponse(success=False, error="禁止访问工作区外的目录")

        # 构建文件路径
        target_file = target_dir / request.file_name

        # 如果文件已存在，添加数字后缀
        if target_file.exists():
            base_name = target_file.stem
            ext = target_file.suffix
            counter = 1
            while target_file.exists():
                target_file = target_dir / f"{base_name}_{counter}{ext}"
                counter += 1

        # 创建文件
        target_file.write_text(request.content, encoding="utf-8")

        # 计算文件 ID（相对路径）
        relative_path = target_file.relative_to(Path(workspace_dir))
        file_id = str(relative_path).replace(os.sep, "/")

        logger.info(f"文件创建成功: {file_id}")

        return CreateFileResponse(
            success=True,
            file_id=file_id,
            file_name=target_file.name,
        )

    except Exception as e:
        logger.error(f"文件创建失败: {e}")
        return CreateFileResponse(success=False, error=str(e))


@router.post("/create-folder", response_model=CreateFolderResponse)
async def create_folder(request: CreateFolderRequest):
    """创建新文件夹

    Args:
        request: 创建文件夹请求，包含文件夹名和可选的父文件夹路径
    """
    try:
        if not request.folder_name:
            return CreateFolderResponse(success=False, error="文件夹名不能为空")

        workspace_dir = deps.config.workspace_dir

        # 构建目标文件夹路径
        if request.parent_folder:
            target_dir = Path(workspace_dir) / request.parent_folder
        else:
            target_dir = Path(workspace_dir)

        # 确保父目录存在
        target_dir.mkdir(parents=True, exist_ok=True)

        # 安全检查：确保目标目录在 workspace 内
        try:
            target_dir.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            return CreateFolderResponse(success=False, error="禁止访问工作区外的目录")

        # 构建文件夹路径
        target_folder = target_dir / request.folder_name

        # 如果文件夹已存在，添加数字后缀
        if target_folder.exists():
            base_name = request.folder_name
            counter = 1
            while target_folder.exists():
                target_folder = target_dir / f"{base_name}_{counter}"
                counter += 1

        # 创建文件夹
        target_folder.mkdir(parents=True, exist_ok=True)

        # 计算文件夹 ID（相对路径）
        relative_path = target_folder.relative_to(Path(workspace_dir))
        folder_id = str(relative_path).replace(os.sep, "/")

        logger.info(f"文件夹创建成功: {folder_id}")

        return CreateFolderResponse(
            success=True,
            folder_id=folder_id,
            folder_name=target_folder.name,
        )

    except Exception as e:
        logger.error(f"文件夹创建失败: {e}")
        return CreateFolderResponse(success=False, error=str(e))


@router.get("/files", response_model=WorkspaceFilesResponse)
async def get_workspace_files():
    """获取工作区文件列表"""
    try:
        workspace_dir = deps.config.workspace_dir

        # 确保目录存在
        if not os.path.exists(workspace_dir):
            logger.warning(f"工作区目录不存在: {workspace_dir}")
            return WorkspaceFilesResponse(files=[], total=0)

        # 构建文件树
        files = build_file_tree(workspace_dir)
        total = count_files(files)

        logger.debug(f"获取工作区文件列表成功，共 {total} 个文件")
        return WorkspaceFilesResponse(files=files, total=total)

    except Exception as e:
        logger.error(f"获取工作区文件列表失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


@router.get("/file", response_model=FileContentResponse)
async def get_file_content(file_id: str = Query(..., description="文件ID（相对路径）")):
    """获取文件内容

    Args:
        file_id: 文件ID，即相对于workspace目录的路径
    """
    try:
        workspace_dir = deps.config.workspace_dir

        # 构建文件完整路径
        file_path = Path(workspace_dir) / file_id

        # 安全检查：确保路径在workspace目录内
        try:
            file_path.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            raise HTTPException(status_code=403, detail="禁止访问工作区外的文件") from None

        # 检查文件是否存在
        if not file_path.exists():
            raise HTTPException(status_code=404, detail=f"文件不存在: {file_id}")

        # 检查是否是文件
        if not file_path.is_file():
            raise HTTPException(status_code=400, detail=f"不是文件: {file_id}")

        # 读取文件内容
        try:
            content = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            # 如果无法以UTF-8读取，尝试以二进制读取并返回base64
            import base64

            content = base64.b64encode(file_path.read_bytes()).decode("ascii")

        return FileContentResponse(
            id=file_id,
            name=file_path.name,
            content=content,
            type="file",
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取文件内容失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


def _validate_rename_request(
    request: RenameFileRequest, old_path: Path, workspace_dir: str
) -> str | None:
    """验证重命名请求，返回错误信息或 None"""
    # 安全检查：确保路径在 workspace 目录内
    try:
        old_path.resolve().relative_to(Path(workspace_dir).resolve())
    except ValueError:
        return "禁止访问工作区外的文件"

    # 检查文件是否存在
    if not old_path.exists():
        return f"文件不存在: {request.file_id}"

    # 验证新文件名
    new_name = request.new_name.strip()
    if not new_name:
        return "文件名不能为空"

    # 检查新文件名是否包含非法字符
    invalid_chars = ["/", "\\", ":", "*", "?", '"', "<", ">", "|"]
    if any(char in new_name for char in invalid_chars):
        return f"文件名包含非法字符: {invalid_chars}"

    # 构建新路径并检查是否已存在
    new_path = old_path.parent / new_name
    if new_path.exists() and new_path != old_path:
        return f"文件已存在: {new_name}"

    return None


@router.post("/rename", response_model=RenameFileResponse)
async def rename_file(request: RenameFileRequest):
    """重命名文件或文件夹

    Args:
        request: 重命名请求，包含文件 ID 和新名称
    """
    try:
        workspace_dir = deps.config.workspace_dir
        old_path = Path(workspace_dir) / request.file_id

        # 验证请求
        error = _validate_rename_request(request, old_path, workspace_dir)
        if error:
            return RenameFileResponse(
                success=False,
                old_id=request.file_id,
                new_id=request.file_id,
                new_name=request.new_name,
                error=error,
            )

        # 执行重命名
        new_name = request.new_name.strip()
        new_path = old_path.parent / new_name
        old_path.rename(new_path)

        # 计算新的文件 ID
        new_relative_path = new_path.relative_to(Path(workspace_dir))
        new_id = str(new_relative_path).replace(os.sep, "/")

        logger.info(f"文件重命名成功: {request.file_id} -> {new_id}")

        return RenameFileResponse(
            success=True,
            old_id=request.file_id,
            new_id=new_id,
            new_name=new_name,
        )

    except Exception as e:
        logger.error(f"重命名文件失败: {e}")
        return RenameFileResponse(
            success=False,
            old_id=request.file_id,
            new_id=request.file_id,
            new_name=request.new_name,
            error=str(e),
        )


@router.post("/save", response_model=SaveFileResponse)
async def save_file(request: SaveFileRequest):
    """保存文件内容

    Args:
        request: 保存文件请求，包含文件ID和内容
    """
    try:
        workspace_dir = deps.config.workspace_dir
        file_path = Path(workspace_dir) / request.file_id

        # 检查文件是否存在
        if not file_path.exists():
            return SaveFileResponse(
                success=False,
                file_id=request.file_id,
                error="文件不存在",
            )

        # 检查是否为文件（不是目录）
        if not file_path.is_file():
            return SaveFileResponse(
                success=False,
                file_id=request.file_id,
                error="路径不是文件",
            )

        # 写入文件内容
        file_path.write_text(request.content, encoding="utf-8")

        # 获取更新时间
        from datetime import datetime

        updated_at = datetime.now().isoformat()

        logger.info(f"保存文件成功: {request.file_id}")
        return SaveFileResponse(
            success=True,
            file_id=request.file_id,
            updated_at=updated_at,
        )

    except Exception as e:
        logger.error(f"保存文件失败: {e}")
        return SaveFileResponse(
            success=False,
            file_id=request.file_id,
            error=str(e),
        )


@router.post("/delete", response_model=DeleteFileResponse)
async def delete_file(request: DeleteFileRequest):
    """删除文件或文件夹

    Args:
        request: 删除请求，包含文件/文件夹 ID
    """
    import shutil

    try:
        if not request.file_id:
            return DeleteFileResponse(
                success=False,
                file_id=request.file_id,
                error="文件 ID 不能为空",
            )

        workspace_dir = deps.config.workspace_dir
        target_path = Path(workspace_dir) / request.file_id

        # 安全检查：确保路径在 workspace 目录内
        try:
            target_path.resolve().relative_to(Path(workspace_dir).resolve())
        except ValueError:
            return DeleteFileResponse(
                success=False,
                file_id=request.file_id,
                error="禁止访问工作区外的文件",
            )

        # 检查文件/文件夹是否存在
        if not target_path.exists():
            return DeleteFileResponse(
                success=False,
                file_id=request.file_id,
                error=f"文件或文件夹不存在: {request.file_id}",
            )

        # 检查是否为受保护的文件（outline.md 在项目根目录）
        # file_id 格式：project_id/outline.md（两层路径表示项目根目录下的文件）
        path_parts = request.file_id.split("/")
        is_root_level = len(path_parts) == 2  # noqa: PLR2004
        if is_root_level and path_parts[1].lower() == "outline.md" and target_path.is_file():
            return DeleteFileResponse(
                success=False,
                file_id=request.file_id,
                error="该文件受保护，无法删除",
            )

        # 删除文件或文件夹
        if target_path.is_file():
            target_path.unlink()
            logger.info(f"文件删除成功: {request.file_id}")
        else:
            # 递归删除文件夹
            shutil.rmtree(target_path)
            logger.info(f"文件夹删除成功: {request.file_id}")

        return DeleteFileResponse(
            success=True,
            file_id=request.file_id,
        )

    except Exception as e:
        logger.error(f"删除失败: {e}")
        return DeleteFileResponse(
            success=False,
            file_id=request.file_id,
            error=str(e),
        )


@router.post("/ai/process", response_model=DocumentAIResponse)
async def process_document_ai(request: DocumentAIRequest):
    """处理文档 AI 操作（非流式）

    Args:
        request: 文档 AI 操作请求
    """
    try:
        # 检查 LLM 是否可用
        if not deps.rag_service.llm_client.is_available():
            return DocumentAIResponse(
                success=False,
                response="",
                action=request.action.value,
                document_name=request.document_name,
                error="LLM服务不可用，请先配置 LLM",
            )

        # 获取对应的系统提示词
        system_prompt = _get_system_prompt(request.action, request.document_name)

        # 构建用户消息
        user_message = _build_user_message(request)

        # 调用 LLM
        response = deps.rag_service.llm_client.client.chat.completions.create(
            model=deps.rag_service.llm_client.model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.7,
        )

        # 记录 token 使用量
        if hasattr(response, "usage") and response.usage:
            log_token_usage(
                model=deps.rag_service.llm_client.model,
                input_tokens=response.usage.prompt_tokens,
                output_tokens=response.usage.completion_tokens,
                endpoint="workspace_ai_process",
                user_query=request.action.value,
                response_type="document_ai",
                feature_type="workspace_assistant",
                additional_info={
                    "action": request.action.value,
                    "document_name": request.document_name,
                    "document_length": len(request.document_content),
                },
            )

        result = response.choices[0].message.content.strip()

        return DocumentAIResponse(
            success=True,
            response=result,
            action=request.action.value,
            document_name=request.document_name,
        )

    except Exception as e:
        logger.error(f"文档 AI 处理失败: {e}")
        return DocumentAIResponse(
            success=False,
            response="",
            action=request.action.value,
            document_name=request.document_name,
            error=str(e),
        )


def _create_token_generator(request: DocumentAIRequest, messages: list[dict]):
    """创建流式 token 生成器"""

    def token_generator():
        try:
            response = deps.rag_service.llm_client.client.chat.completions.create(
                model=deps.rag_service.llm_client.model,
                messages=messages,
                temperature=0.7,
                stream=True,
                stream_options={"include_usage": True},
            )

            total_content = ""
            usage_info = None

            for chunk in response:
                if hasattr(chunk, "usage") and chunk.usage:
                    usage_info = chunk.usage

                if chunk.choices and len(chunk.choices) > 0 and chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    total_content += content
                    yield content

            _log_stream_usage(usage_info, request, total_content)

        except Exception as e:
            logger.error(f"流式生成失败: {e}")
            yield f"\n[错误] 流式生成失败: {e}"

    return token_generator


def _log_stream_usage(usage_info, request: DocumentAIRequest, total_content: str):
    """记录流式响应的 token 使用量"""
    if not usage_info:
        return

    try:
        log_token_usage(
            model=deps.rag_service.llm_client.model,
            input_tokens=usage_info.prompt_tokens,
            output_tokens=usage_info.completion_tokens,
            endpoint="workspace_ai_stream",
            user_query=request.action.value,
            response_type="stream",
            feature_type="workspace_assistant",
            additional_info={
                "action": request.action.value,
                "document_name": request.document_name,
                "document_length": len(request.document_content),
                "response_length": len(total_content),
            },
        )
    except Exception as log_error:
        logger.error(f"记录 token 使用量失败: {log_error}")


@router.post("/ai/stream")
async def process_document_ai_stream(request: DocumentAIRequest):
    """处理文档 AI 操作（流式输出）

    Args:
        request: 文档 AI 操作请求
    """
    try:
        if not deps.rag_service.llm_client.is_available():

            async def error_generator():
                yield "LLM服务不可用，请先配置 LLM"

            return StreamingResponse(error_generator(), media_type="text/plain; charset=utf-8")

        system_prompt = _get_system_prompt(request.action, request.document_name)
        user_message = _build_user_message(request)
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ]

        token_generator = _create_token_generator(request, messages)
        headers = {"Cache-Control": "no-cache", "X-Accel-Buffering": "no"}

        return StreamingResponse(
            token_generator(), media_type="text/plain; charset=utf-8", headers=headers
        )

    except Exception as e:
        logger.error(f"文档 AI 流式处理失败: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


def _get_system_prompt(action: DocumentAction, document_name: str | None = None) -> str:
    """获取对应操作的系统提示词

    Args:
        action: 操作类型
        document_name: 文档名称

    Returns:
        系统提示词
    """
    # 使用字典映射替代多个 if-elif 语句
    action_prompt_map = {
        DocumentAction.SUMMARIZE: "summarize",
        DocumentAction.IMPROVE: "improve",
        DocumentAction.EXPLAIN: "explain",
        DocumentAction.BEAUTIFY: "beautify",
        DocumentAction.EXPAND: "expand",
        DocumentAction.CONDENSE: "condense",
        DocumentAction.CORRECT: "correct",
        DocumentAction.TRANSLATE: "translate",
    }

    # 特殊处理 CUSTOM 操作
    if action == DocumentAction.CUSTOM:
        prompt = get_prompt("workspace_assistant", "custom_chat")
        return prompt.replace("{document_name}", document_name or "未命名文档")

    prompt_key = action_prompt_map.get(action, "summarize")
    return get_prompt("workspace_assistant", prompt_key)


def _build_user_message(request: DocumentAIRequest) -> str:
    """构建用户消息

    Args:
        request: 文档 AI 操作请求

    Returns:
        用户消息
    """
    # 文本编辑操作（美化、扩写、缩写、修正、翻译）只需要发送选中的文本
    text_edit_actions = [
        DocumentAction.BEAUTIFY,
        DocumentAction.EXPAND,
        DocumentAction.CONDENSE,
        DocumentAction.CORRECT,
        DocumentAction.TRANSLATE,
    ]

    if request.action in text_edit_actions:
        # 对于文本编辑操作，直接发送文本内容
        return request.document_content
    elif request.action == DocumentAction.CUSTOM and request.custom_prompt:
        # 自定义对话模式
        return f"""**文档内容：**
```
{request.document_content}
```

**用户问题：**
{request.custom_prompt}"""
    else:
        # 预定义操作模式（总结、改进、解释）
        doc_name_str = f"**文件名：** {request.document_name}\n\n" if request.document_name else ""
        return f"""{doc_name_str}**文档内容：**
```
{request.document_content}
```"""
